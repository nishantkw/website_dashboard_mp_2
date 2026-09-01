import json
import re
from pathlib import Path
from docx import Document

DOCX = Path(r"c:/Users/bk455/Documents/website_dashboard/database_documentation.docx")
PARTIAL = Path(r"c:/Users/bk455/Documents/website_dashboard/backend/src/data/doc_tables_extract.json")
OUT = Path(r"c:/Users/bk455/Documents/website_dashboard/backend/src/data/doc_schema_full.json")

EXPECTED = [
    "claim_paid_t",
    "claim_paid_t_portability",
    "hospital_master_with_quality_certification",
    "hospital_master_with_quality_certification_final",
    "icd_data_doctor_details",
    "lms_user_course_completion_status",
    "m_lookup",
    "m_source_data",
    "payment_dtls",
    "pro_workflow_users_t",
    "t_bis_beneficiary_disabled",
    "t_bis_beneficiary_dtls",
    "t_card_printing_status",
    "t_deempanelment_details",
    "t_hem_hospital",
    "t_hem_sch_mapping",
    "t_patient_dtl",
    "t_patient_dtls",
    "t_morth_patient_details",
    "t_payment_dtls",
    "t_suspicious_api_case_data",
    "t_suspicious_api_case_dtls",
    "t_workflow_transaction_audit",
    "t_workflow_transaction_audit_hem",
    "tms_recovery",
    "treatment_dtls",
    "ump_users_v",
    "user_master_ump",
    "workflow_users_t",
]

TYPE_ALIASES = {
    "character varying": "varchar",
    "varchar": "varchar",
    "text": "text",
    "integer": "integer",
    "int": "integer",
    "bigint": "bigint",
    "smallint": "smallint",
    "numeric": "numeric",
    "decimal": "numeric",
    "double precision": "double precision",
    "real": "real",
    "boolean": "boolean",
    "date": "date",
    "timestamp without time zone": "timestamp",
    "timestamp with time zone": "timestamptz",
    "timestamp": "timestamp",
    "time without time zone": "time",
}


def normalize_type(raw: str) -> str:
    t = (raw or "").strip().lower()
    t = re.sub(r"\s+", " ", t)
    if not t:
        return "unknown"
    return TYPE_ALIASES.get(t, t)


def cell_text(cell) -> str:
    return re.sub(r"\s+", " ", cell.text.replace("\n", " ")).strip()


def extract_columns_from_table(tbl):
    if not tbl.rows:
        return []
    hdr = [cell_text(c).lower() for c in tbl.rows[0].cells]
    name_idx = None
    type_idx = None
    for i, h in enumerate(hdr):
        if "column name" in h or h == "column":
            name_idx = i
        if "data type" in h or h == "type":
            type_idx = i
    if name_idx is None:
        name_idx = 1 if len(hdr) > 1 else 0
    if type_idx is None:
        type_idx = 2 if len(hdr) > 2 else 1
    cols = []
    for row in tbl.rows[1:]:
        cells = row.cells
        if name_idx >= len(cells):
            continue
        name = cell_text(cells[name_idx])
        if not name or name == "#":
            continue
        dtype = normalize_type(cell_text(cells[type_idx]) if type_idx < len(cells) else "")
        cols.append({"name": name, "type": dtype})
    return cols


def parse_table_names_from_paragraphs(doc):
    names = []
    for p in doc.paragraphs:
        m = re.match(r"^5\.\d+\.\s*Table:\s*(.+)$", p.text.strip())
        if m:
            names.append(m.group(1).strip())
    return names


def extract_from_paragraph_patterns(doc, table_name):
    """Fallback: lines like 'column_name (type)' in section text."""
    cols = []
    in_section = False
    section_re = re.compile(rf"^5\.\d+\.\s*Table:\s*{re.escape(table_name)}\s*$", re.I)
    next_section = re.compile(r"^5\.\d+\.\s*Table:\s*", re.I)
    col_line = re.compile(
        r"^([a-z_][a-z0-9_]*)\s*[\(\[]?\s*(character varying|varchar|text|integer|bigint|numeric|timestamp|date|boolean|double precision|real|smallint)",
        re.I,
    )
    for p in doc.paragraphs:
        t = p.text.strip()
        if section_re.match(t):
            in_section = True
            continue
        if in_section and next_section.match(t) and table_name.lower() not in t.lower():
            break
        if in_section:
            m = col_line.match(t)
            if m:
                cols.append({"name": m.group(1), "type": normalize_type(m.group(2))})
    return cols


def load_partial():
    if not PARTIAL.exists():
        return {}
    data = json.loads(PARTIAL.read_text(encoding="utf-8"))
    out = {}
    for k, v in data.items():
        if isinstance(v, dict) and "columns" in v:
            names = v["columns"]
            out[k] = [{"name": n, "type": "unknown"} for n in names if isinstance(n, str)]
        elif isinstance(v, list):
            out[k] = v
    return out


def main():
    doc = Document(str(DOCX))
    heading_names = parse_table_names_from_paragraphs(doc)

    column_tables = []
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        hdr = [cell_text(c).lower() for c in tbl.rows[0].cells]
        if any("column name" in h for h in hdr) and any("data type" in h for h in hdr):
            column_tables.append(tbl)

    partial = load_partial()
    result = {}
    meta = {"source": str(DOCX), "schema": "dmart_mp", "tables": {}}

    for i, tname in enumerate(EXPECTED):
        cols = []
        source = None
        if i < len(column_tables):
            cols = extract_columns_from_table(column_tables[i])
            source = "docx_table"
        if not cols and i < len(heading_names) and heading_names[i] == tname:
            if i < len(column_tables):
                cols = extract_columns_from_table(column_tables[i])
                source = "docx_table"
        if not cols:
            cols = extract_from_paragraph_patterns(doc, tname)
            if cols:
                source = "paragraph_pattern"
        if not cols and tname in partial:
            cols = partial[tname]
            source = "partial_json"
        result[tname] = cols
        meta["tables"][tname] = {
            "column_count": len(cols),
            "source": source,
            "heading_match": heading_names[i] if i < len(heading_names) else None,
        }

    OUT.parent.mkdir(parents=True, exist_ok=True)
    payload = {"schema": "dmart_mp", "tables": result, "_meta": meta}
    OUT.write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    # summary
    missing = [t for t, c in result.items() if not c]
    print("written", OUT)
    print("tables", len(result), "with columns", sum(1 for c in result.values() if c))
    print("total columns", sum(len(c) for c in result.values()))
    if missing:
        print("missing", missing)
    for t in EXPECTED[:3]:
        print(t, len(result[t]), result[t][0] if result[t] else None)


if __name__ == "__main__":
    main()
